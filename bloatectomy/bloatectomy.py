#!/usr/bin/env python
# -*- coding: UTF-8 -*-

#Bloatectomy: a method for the identification and removal of duplicate text in the bloated notes of electronic health records and other documents.
#Copyright (C) 2020  Summer K. Rankin, Roselie A. Bright, Katherine R. Dowdy
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU General Public License as published by
#    the Free Software Foundation, either version 3 of the License, or
#    (at your option) any later version.

#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU General Public License for more details.

#    You should have received a copy of the GNU General Public License
#    along with this program.  If not, see <https://www.gnu.org/licenses/>.

#Authors: Summer K. Rankin summerkrankin@gmail.com, Roselie A. Bright roselie.bright@fda.hhs.gov, Katherine R. Dowdy katerdowdy@gmail.com

import re
import sys

class bloatectomy():
    def __init__(self, input_text,  path = '', filename='bloatectomized_file',
                 display=False, style='highlight', output='html', output_numbered_tokens=False, output_original_tokens=False,
                 regex1=r"(.+?\.[\s\n]+)", regex2=r"(?=\n\s*[A-Z1-9#-]+.*)", postgres_engine=None, postgres_table=None):

        self.path = path
        self.filename = filename
        self.display = display
        self.style = style # =['highlight','bold','remove']
        self.output_numbered_tokens  = output_numbered_tokens
        self.output_original_numbered_tokens  = output_original_tokens
        self.output = output
        self.regex1 = regex1
        self.regex2 = regex2
        self.postgres_table = postgres_table
        self.engine  = postgres_engine

        assert float(sys.version[0:3]) >= 3.7, "Must use python 3.7.0 or higher for the regular expressions to work correctly."

        try:
            if  input_text.split('.')[1] == 'docx' or input_text.split('.')[1] == 'doc':
                import docx
                doc = docx.Document(input_text)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                self.input_text = '\n'.join(fullText)
                print(style + "ing duplications in " + input_text + ". Output file = " + path + filename + '.' + output)
                bloatectomy.main(self)

            elif input_text.split('.')[1] == 'txt' or input_text.split('.')[1] == 'rtf':
                with open(input_text) as file:
                    self.input_text = file.read()
                print(style + "ing duplications in " + input_text + ". Output file = " + path + filename + '.' + output)
                bloatectomy.main(self)

            else:
                assert type(input_text) == str, "unsupported format"
                self.input_text = input_text
                print(style + "ing duplications. Output file = " + path + filename + '.' + output)
                bloatectomy.main(self)

        except IndexError:
            assert type(input_text) == str, "unsupported format"
            self.input_text = input_text
            print(style + "ing duplications. Output file = " + path + filename + '.' + output)
            bloatectomy.main(self)

        except AttributeError:
            import numpy
            import pandas as pd
            assert (type(input_text) == list or type(input_text) == numpy.ndarray), "unsupported format"
            print('pulling notes from postgres database')
            # select one of the hadm_id s in the list and get the concatenated notes for that hadm_id
            for i in input_text:
                print(style + "ing duplications in ID " + str(i))
                # in this table, we have concatenated all notes from each hadm_id into a single document
                query = """SELECT text FROM {0} WHERE hadm_id IN ({1})"""
                query = query.format(self.postgres_table, i)
                pt_text = pd.read_sql(query, self.engine)
                self.input_text = ''
                self.input_text = pt_text.text.to_list()[0]
                self.filename = ''
                self.filename = filename + '_' + str(i)
                bloatectomy.main(self)
                print("Output file = " + path + self.filename + '.' + output)

    def main(self):
        bloatectomy.tokenize_mark(self)
        if self.output=='html':
            bloatectomy.make_html(self)
        else:
            bloatectomy.make_docx(self)

    def make_html(self):
        """Takes the output of the just_replication_detection (list of strings) and returns an html file (in path + filename) for the admission with duplicates highlighted"""

        file_name =  str(self.path) + str(self.filename) + '.html'
        uniq = str("\n ".join(self.tokens))
        # replace line feed characters with html linebreaks
        uniq = uniq.replace("\n", "<br>")
        # save bloatectomized file as an html
        with open(file_name, "w") as file:
            file.write(uniq)

        if self.output_numbered_tokens == True:
            with open(str(self.path) + str(self.filename) + '_token_numbers.txt',"w") as file:
                for i in self.numbered_tokens:
                 file.write(str("$ ".join(i) + '\n'))

        if self.output_original_numbered_tokens == True:
            with open(str(self.path) + str(self.filename) + '_original_token_numbers.txt',"w") as file:
                for i in self.original_numbered_tokens:
                 file.write(str("$ ".join(i) + '\n'))

        if self.display==True:
            print(uniq)

    def make_docx(self):
        """Takes the output of the just_replication_detection (list of strings) and returns a docx file (in path + filename) for the admission with duplicates highlighted"""
        import docx
        document = docx.Document()
        p = document.add_paragraph()
        run = p.add_run()
        font = run.font
        file_name =  str(self.path) + str(self.filename) + '.docx'
        # adding the first entry, not duplicate by nature
        p.add_run(self.tokens[0])
        # iterating through all the tokens in text
        for i in range(1, len(self.tokens)):
            if bool(re.search('[<][m][a][r][k][>]', self.tokens[i])):
                new_text = self.tokens[i].replace("<mark>", "").replace("</mark>", "")
                run = p.add_run()
                run.add_text(new_text)
                run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                run.add_break(docx.enum.text.WD_BREAK.TEXT_WRAPPING)
            elif bool(re.search('[<][b][>]', self.tokens[i])):
                new_text = self.tokens[i].replace("<b>", "").replace("</b>", "")
                run = p.add_run()
                run.add_text(new_text)
                run.font.bold = True
                run.add_break(docx.enum.text.WD_BREAK.TEXT_WRAPPING)
            else:
                run = p.add_run()
                run.add_text(self.tokens[i])
                run.add_break(docx.enum.text.WD_BREAK.TEXT_WRAPPING)

        document.save(file_name)
        if self.display==True:
            print(self.tokens)
        if self.output_numbered_tokens == True:
            with open(str(self.path) + str(self.filename) + '_token_numbers.txt',"w") as file:
                for i in self.numbered_tokens:
                    file.write(str("$ ".join(i) + '\n'))
        if self.output_original_numbered_tokens == True:
            with open(str(self.path) + str(self.filename) + '_original_token_numbers.txt',"w") as file:
                for i in self.original_numbered_tokens:
                 file.write(str("$ ".join(i) + '\n'))

    def tokenize2(regex, token_in):
        """
        Tokenize (2nd time) on  a line feed character.
        1. for each token, split if a line feed character is followed by
        2. a capital letter, or a dash, or a number
        """
        tok_new = []
        # find any \n followed by an uppercase letter, a number, or a dash
        sent_token =re.split(regex, token_in)
        # replace \n with a space with a space
        sent_token = [re.sub(r"$\n+","",i) for i in sent_token] # remove from end
        sent_token = [re.sub(r"^\n", "", i) for i in sent_token] #remove from front
            # line feeds + whitespace or not
        sent_token = [re.sub(r"\s+\n\s+", " ", i) for i in sent_token]
        sent_token = [re.sub(r"\s+\n", " ", i) for i in sent_token]
        sent_token = [re.sub(r"\n\s+", " ", i) for i in sent_token]
        sent_token = [re.sub(r"\n", " ", i) for i in sent_token]
        #remove front/end whitespace
        sent_token = [i.strip(' ') for i in sent_token]
        for i in sent_token:
            if i != '':
                tok_new.append(i)
        return tok_new

    def number_tokens(token):
        """create a list of enumerated (numbered) tokens"""
        for enum_num, enum_token in enumerate(token):
            yield str(enum_num), enum_token

    def tokenize_mark(self):
        """
        1. Take in raw text and do initial tokenization: on periods followed by one or more space, tab, or line feed character.
        2. Secondary tokenization of each token on line feed character followed by a capital letter, or a number, or a dash.
        3. Add tags to or remove duplicate tokens.
        """
        # tokenize 1
        tok = re.split(self.regex1, self.input_text, flags=re.DOTALL)
        # whitespace around tokens can cause a duplicate to be missed
        tok = [i.strip(' ') for i in tok]
        #tokenize 2
        new_tok = []
        for num, t in enumerate(tok):
            n_tok = bloatectomy.tokenize2(self.regex2, t)
            new_tok.extend(n_tok)
        # save original data as numbered list
        self.original_numbered_tokens = []
        self.original_numbered_tokens = list(bloatectomy.number_tokens(new_tok))
        # detect and mark/remove duplicates
        self.tokens = []
        self.tokens = list(bloatectomy.mark_duplicates(self, new_tok))
        # save bloatectomized tokens as a numbered list
        self.numbered_tokens = []
        self.numbered_tokens = list(bloatectomy.number_tokens(self.tokens))

    def mark_duplicates(self, input_tokens):
        '''
        Function uses a set() and list to generate each token with tags (of selected style) added to duplicate tokens.
            INPUT: input_tokens = string of tokenized text (can be sentences, paragraphs, words etc)
                   style = ['bold','higlight','remov'] what to do with duplicate text.
            OUTPUT: yield a single token at a time (generator) until the end of the input_tokens.
    '''
        if self.style == 'bold':
            tag = '<b>'
            tag_end = '</b>'
            remov = False
        elif self.style == 'highlight':
            tag = '<mark>'
            tag_end = '</mark>'
            remov = False
        elif self.style == 'remov':
            remov = True
        else:
            print("Select a style for duplicate text: 'bold' or 'highlight")
        # create hash of tokens
        tokens_set = set()
        tokens_set_add = tokens_set.add
        for token in input_tokens:
            #skip any empty tokens
            if token == '':
                pass
            elif token not in tokens_set:
                tokens_set_add(token)
                yield token
            elif remov == False:
                yield tag + token + tag_end
            elif remov == True:
                pass
